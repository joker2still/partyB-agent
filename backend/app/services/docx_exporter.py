from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


EXPORTS_DIR = Path(__file__).resolve().parent.parent.parent / "exports"
SECTION_TITLE_MAP = {
    "\u6559\u80b2\u80cc\u666f": "education",
    "\u6559\u80b2\u7ecf\u5386": "education",
    "\u4e13\u4e1a\u6280\u80fd": "skills",
    "\u6280\u80fd": "skills",
    "\u9879\u76ee\u5b9e\u8df5": "project_experience",
    "\u9879\u76ee\u7ecf\u5386": "project_experience",
    "\u5de5\u4f5c\u5b9e\u8df5\u4e0e\u7ecf\u5386": "work_experience",
    "\u5de5\u4f5c\u7ecf\u5386": "work_experience",
    "\u81ea\u6211\u8bc4\u4ef7": "self_evaluation",
    "\u6c42\u804c\u610f\u5411": "target_position",
}


def _sanitize_session_id(session_id: str) -> str:
    cleaned = "".join(char for char in session_id if char.isalnum() or char in {"-", "_"})
    return cleaned or "default"


def _set_run_font(run, size: int, bold: bool = False) -> None:
    run.bold = bold
    font = run.font
    font.size = Pt(size)
    font.name = "Arial"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_paragraph_bottom_border(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)
    paragraph_format.space_after = Pt(8)

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_tag = qn(f"w:{edge}")
        element = tc_borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _extract_sections_from_draft(markdown_text: str) -> dict[str, str]:
    parsed = {
        "name": "",
        "target_position": "",
        "education": "",
        "skills": "",
        "project_experience": "",
        "work_experience": "",
        "self_evaluation": "",
    }
    current_key = ""
    buffers: dict[str, list[str]] = {key: [] for key in parsed if key != "name"}

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            heading = line[2:].strip()
            if not parsed["name"]:
                parsed["name"] = heading
            current_key = SECTION_TITLE_MAP.get(heading, "")
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            current_key = SECTION_TITLE_MAP.get(heading, "")
            continue

        if current_key and current_key in buffers:
            buffers[current_key].append(line)

    for key, lines in buffers.items():
        parsed[key] = "\n".join(lines).strip()
    return parsed


def _pick_value(slots: dict, draft_data: dict, key: str, default: str = "") -> str:
    value = slots.get(key, "")
    if isinstance(value, str) and value.strip():
        return value.strip()
    draft_value = draft_data.get(key, "")
    if isinstance(draft_value, str) and draft_value.strip():
        return draft_value.strip()
    return default


def _add_lines(document: Document, value: str) -> None:
    if not value.strip():
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\u5f85\u8865\u5145")
        _set_run_font(run, 10)
        paragraph.paragraph_format.space_after = Pt(4)
        return

    for raw_line in value.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("- ") or line.startswith("\u00b7 "):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:].strip())
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)

        _set_run_font(run, 10)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25


def _add_section(document: Document, title: str, content: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(title)
    _set_run_font(run, 12, bold=True)
    _set_paragraph_bottom_border(heading)
    _add_lines(document, content)


def _build_header(document: Document, slots: dict, draft_data: dict) -> None:
    header_table = document.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_table.autofit = False

    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    left_cell.width = Cm(13.5)
    right_cell.width = Cm(3.3)

    placeholder = "\u5f85\u8865\u5145"
    name = _pick_value(slots, draft_data, "name", placeholder)
    info_values = [
        f"\u624b\u673a\uff1a{_pick_value(slots, draft_data, 'phone', placeholder)}",
        f"\u90ae\u7bb1\uff1a{_pick_value(slots, draft_data, 'email', placeholder)}",
        f"\u5e74\u9f84\uff1a{_pick_value(slots, draft_data, 'age', placeholder)}",
        f"\u73b0\u5c45\uff1a{_pick_value(slots, draft_data, 'location', placeholder)}",
        f"\u6c42\u804c\u610f\u5411\uff1a{_pick_value(slots, draft_data, 'target_position', placeholder)}",
    ]

    name_paragraph = left_cell.paragraphs[0]
    name_paragraph.paragraph_format.space_after = Pt(8)
    name_run = name_paragraph.add_run(name)
    _set_run_font(name_run, 18, bold=True)

    info_paragraph = left_cell.add_paragraph()
    info_paragraph.paragraph_format.line_spacing = 1.25
    info_paragraph.paragraph_format.space_after = Pt(2)
    info_run = info_paragraph.add_run(" / ".join(info_values))
    _set_run_font(info_run, 9)

    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_border(right_cell)
    photo_paragraph = right_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_paragraph.paragraph_format.space_before = Pt(22)
    photo_paragraph.paragraph_format.space_after = Pt(22)
    photo_run = photo_paragraph.add_run("\u7167\u7247")
    _set_run_font(photo_run, 11, bold=True)


def export_structured_resume_docx(slots: dict, session_id: str) -> str:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

    filename = f"resume_{_sanitize_session_id(session_id)}.docx"
    output_path = EXPORTS_DIR / filename
    safe_slots = dict(slots) if isinstance(slots, dict) else {}
    markdown_text = safe_slots.get("resume_draft", "")
    if not isinstance(markdown_text, str):
        markdown_text = ""

    draft_data = _extract_sections_from_draft(markdown_text)

    document = Document()
    _set_document_defaults(document)
    _build_header(document, safe_slots, draft_data)
    document.add_paragraph()

    _add_section(document, "\u6559\u80b2\u80cc\u666f", _pick_value(safe_slots, draft_data, "education"))
    _add_section(document, "\u4e13\u4e1a\u6280\u80fd", _pick_value(safe_slots, draft_data, "skills"))
    _add_section(document, "\u9879\u76ee\u5b9e\u8df5", _pick_value(safe_slots, draft_data, "project_experience"))
    _add_section(
        document,
        "\u5de5\u4f5c\u5b9e\u8df5\u4e0e\u7ecf\u5386",
        _pick_value(safe_slots, draft_data, "work_experience"),
    )
    _add_section(document, "\u81ea\u6211\u8bc4\u4ef7", _pick_value(safe_slots, draft_data, "self_evaluation"))

    document.save(output_path)
    return str(output_path)


def export_resume_docx(markdown_text: str, session_id: str) -> str:
    return export_structured_resume_docx({"resume_draft": markdown_text}, session_id)
